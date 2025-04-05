import os
import typer
from pdf2image import convert_from_path
from typing import List, Optional
from docx import Document
from docx.shared import Inches
import tempfile


def extract_images_from_pdf(pdf_path, output_folder):
    """Extracts pages from a PDF as images and saves them to an output folder."""
    # Convert PDF pages to images
    images = convert_from_path(pdf_path)

    # Save each page as a separate image file
    for i, image in enumerate(images):
        image_name = f"page_{i + 1}.png"
        image_path = os.path.join(output_folder, image_name)
        image.save(image_path, "PNG")
        print(f"Saved: {image_path}")


def process_pdf_in_folder(folder_path):
    """Processes all PDFs found in the given folder and its subdirectories."""
    # Walk through the folder and subfolders
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            if file_name.lower().endswith(".pdf"):
                # Construct full file path
                pdf_path = os.path.join(root, file_name)

                # Create an output folder named after the PDF (without the .pdf extension)
                output_folder = os.path.join(root, os.path.splitext(file_name)[0])
                os.makedirs(output_folder, exist_ok=True)

                # Extract images from the PDF and save them in the output folder
                print(f"Processing PDF: {pdf_path}")
                extract_images_from_pdf(pdf_path, output_folder)


def extract_images_to_word(pdf_path, output_doc_path=None):
    """Extracts pages from a PDF as images and adds them to a Word document."""
    # Convert PDF pages to images
    images = convert_from_path(pdf_path)

    # Create a new Word document
    doc = Document()

    # If no output path specified, use the PDF filename with .docx extension
    if output_doc_path is None:
        output_doc_path = os.path.splitext(pdf_path)[0] + ".docx"

    # Add a title with the PDF filename
    doc.add_heading(f"Images from {os.path.basename(pdf_path)}", 0)

    # Add each page as an image in the document
    with tempfile.TemporaryDirectory() as temp_dir:
        for i, image in enumerate(images):
            # Save the image temporarily
            temp_image_path = os.path.join(temp_dir, f"temp_page_{i + 1}.png")
            image.save(temp_image_path, "PNG")

            # Add page number as heading
            doc.add_heading(f"Page {i + 1}", 1)

            # Add the image to the document
            doc.add_picture(temp_image_path, width=Inches(6))

            # Add a page break after each image (except the last one)
            if i < len(images) - 1:
                doc.add_page_break()

            print(f"Added page {i + 1} to Word document")

    # Save the document
    doc.save(output_doc_path)
    print(f"Word document saved: {output_doc_path}")
    return output_doc_path


def process_pdf_to_word(folder_path, output_folder=None):
    """Processes all PDFs found in the given folder and converts them to Word documents."""
    documents_created = []

    # Walk through the folder and subfolders
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            if file_name.lower().endswith(".pdf"):
                # Construct full file path
                pdf_path = os.path.join(root, file_name)

                # Determine output path
                if output_folder:
                    os.makedirs(output_folder, exist_ok=True)
                    output_doc_path = os.path.join(
                        output_folder, os.path.splitext(file_name)[0] + ".docx"
                    )
                else:
                    output_doc_path = os.path.splitext(pdf_path)[0] + ".docx"

                # Extract images from the PDF and save to Word doc
                print(f"Processing PDF to Word: {pdf_path}")
                doc_path = extract_images_to_word(pdf_path, output_doc_path)
                documents_created.append(doc_path)

    return documents_created


app = typer.Typer(help="Extract images from PDF files")


@app.command()
def extract_images(
    folders: List[str] = typer.Argument(
        ..., help="Folders containing PDF files to process"
    )
):
    """Process PDFs in the specified folders and extract images from them."""
    for folder in folders:
        if os.path.isdir(folder):
            typer.echo(f"Scanning folder: {folder}")
            process_pdf_in_folder(folder)
        else:
            typer.echo(f"Warning: {folder} is not a valid directory.")


@app.command()
def pdf_to_word(
    folders: List[str] = typer.Argument(
        ..., help="Folders containing PDF files to process"
    ),
    output_folder: Optional[str] = typer.Option(
        None, "--output", "-o", help="Output folder for the Word documents"
    ),
):
    """Convert PDFs to Word documents with each page as an image."""
    documents_created = []

    for folder in folders:
        if os.path.isdir(folder):
            typer.echo(f"Scanning folder: {folder}")
            docs = process_pdf_to_word(folder, output_folder)
            documents_created.extend(docs)
        else:
            typer.echo(f"Warning: {folder} is not a valid directory.")

    if documents_created:
        typer.echo(f"Created {len(documents_created)} Word document(s):")
        for doc in documents_created:
            typer.echo(f"  - {doc}")
    else:
        typer.echo("No Word documents were created.")


if __name__ == "__main__":
    app()
