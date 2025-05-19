import os
import typer
from pdf2image import convert_from_path
from typing import List, Optional
from docx import Document
from docx.shared import Inches
import tempfile
import img2pdf  # New import for image to PDF conversion
import zipfile  # For extracting images from docx
import re  # For matching file patterns
import shutil  # For copying extracted files
from PIL import Image  # For image manipulation


def extract_images_from_pdf(pdf_path, output_folder):
    """Extracts pages from a PDF as images and saves them to an output folder."""
    # Convert PDF pages to images
    images = convert_from_path(pdf_path)

    # Save each page as a separate image file
    for i, image in enumerate(images):
        image_name = f"page_{i + 1}.png"
        image_path = os.path.join(output_folder, image_name)
        image.save(image_path, "PNG")
        print(f"Saved: {image_path}")


def process_pdf_in_folder(folder_path):
    """Processes all PDFs found in the given folder and its subdirectories."""
    # Walk through the folder and subfolders
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            if file_name.lower().endswith(".pdf"):
                # Construct full file path
                pdf_path = os.path.join(root, file_name)

                # Create an output folder named after the PDF (without the .pdf extension)
                output_folder = os.path.join(root, os.path.splitext(file_name)[0])
                os.makedirs(output_folder, exist_ok=True)

                # Extract images from the PDF and save them in the output folder
                print(f"Processing PDF: {pdf_path}")
                extract_images_from_pdf(pdf_path, output_folder)


def extract_images_to_word(pdf_path, output_doc_path=None):
    """Extracts pages from a PDF as images and adds them to a Word document."""
    # Convert PDF pages to images
    images = convert_from_path(pdf_path)

    # Create a new Word document
    doc = Document()

    # If no output path specified, use the PDF filename with .docx extension
    if output_doc_path is None:
        output_doc_path = os.path.splitext(pdf_path)[0] + ".docx"

    # Add a title with the PDF filename
    doc.add_heading(f"Images from {os.path.basename(pdf_path)}", 0)

    # Add each page as an image in the document
    with tempfile.TemporaryDirectory() as temp_dir:
        for i, image in enumerate(images):
            # Save the image temporarily
            temp_image_path = os.path.join(temp_dir, f"temp_page_{i + 1}.png")
            image.save(temp_image_path, "PNG")

            # Add page number as heading
            doc.add_heading(f"Page {i + 1}", 1)

            # Add the image to the document
            doc.add_picture(temp_image_path, width=Inches(6))

            # Add a page break after each image (except the last one)
            if i < len(images) - 1:
                doc.add_page_break()

            print(f"Added page {i + 1} to Word document")

    # Save the document
    doc.save(output_doc_path)
    print(f"Word document saved: {output_doc_path}")
    return output_doc_path


def process_pdf_to_word(folder_path, output_folder=None):
    """Processes all PDFs found in the given folder and converts them to Word documents."""
    documents_created = []

    # Walk through the folder and subfolders
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            if file_name.lower().endswith(".pdf"):
                # Construct full file path
                pdf_path = os.path.join(root, file_name)

                # Determine output path
                if output_folder:
                    os.makedirs(output_folder, exist_ok=True)
                    output_doc_path = os.path.join(
                        output_folder, os.path.splitext(file_name)[0] + ".docx"
                    )
                else:
                    output_doc_path = os.path.splitext(pdf_path)[0] + ".docx"

                # Extract images from the PDF and save to Word doc
                print(f"Processing PDF to Word: {pdf_path}")
                doc_path = extract_images_to_word(pdf_path, output_doc_path)
                documents_created.append(doc_path)

    return documents_created


def find_and_sort_images(folder_path):
    """Find all image files in a folder and sort them by filename."""
    # Common image file extensions
    image_extensions = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff")

    images = []

    # Walk through the folder and find image files
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            if file_name.lower().endswith(image_extensions):
                file_path = os.path.join(root, file_name)
                images.append(file_path)

    # Sort images by filename without extension
    images.sort(key=lambda x: os.path.splitext(os.path.basename(x))[0])

    return images


def create_pdf_from_images(image_paths, output_pdf_path):
    """Create a PDF file from a list of image paths."""
    if not image_paths:
        print("No images found to convert to PDF.")
        return None

    # Create PDF from images
    with open(output_pdf_path, "wb") as f:
        f.write(img2pdf.convert(image_paths))

    print(f"PDF created: {output_pdf_path}")
    return output_pdf_path


def process_images_to_pdf(folder_path, output_pdf=None):
    """Process all images in a folder and create a PDF file."""
    # Find and sort images
    image_paths = find_and_sort_images(folder_path)

    if not image_paths:
        print(f"No images found in folder: {folder_path}")
        return None

    # If no output path specified, use the folder name with .pdf extension
    if output_pdf is None:
        # Get parent directory and use folder name for the PDF
        parent_dir = os.path.dirname(folder_path)
        folder_name = os.path.basename(folder_path)
        output_pdf = os.path.join(parent_dir, folder_name + ".pdf")

    # Create PDF from images
    return create_pdf_from_images(image_paths, output_pdf)


def convert_pdf_to_image_based_pdf(pdf_path):
    """
    Process a PDF file:
    1. Extract its pages as images into a temporary folder
    2. Combine those images into a new PDF with '-images' suffix
    Returns the path to the new PDF if successful
    """
    # Get path components
    pdf_dir = os.path.dirname(pdf_path)
    pdf_name = os.path.basename(pdf_path)
    base_name = os.path.splitext(pdf_name)[0]
    output_pdf = os.path.join(pdf_dir, f"{base_name}-images.pdf")
    
    # Create temporary folder for images
    with tempfile.TemporaryDirectory() as temp_dir:
        # Extract images from PDF
        extract_images_from_pdf(pdf_path, temp_dir)
        
        # Find and sort images
        image_paths = find_and_sort_images(temp_dir)
        
        if not image_paths:
            print(f"No images extracted from {pdf_path}")
            return None
            
        # Create new PDF from the images
        created_pdf = create_pdf_from_images(image_paths, output_pdf)
        
    return created_pdf


def extract_images_from_docx(docx_path, output_folder):
    """
    Extract all images from a Word document (.docx) and save them to the specified folder
    with sequential numbering.
    """
    # Ensure output folder exists
    os.makedirs(output_folder, exist_ok=True)
    
    # A .docx file is actually a zip file, so we can extract its contents
    image_count = 0
    
    # Create a temporary directory to extract the docx contents
    with tempfile.TemporaryDirectory() as temp_dir:
        # Extract the docx contents to the temp directory
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Find all image files in the docx
            image_files = [f for f in docx_zip.namelist() 
                          if f.startswith('word/media/') and 
                          f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
            
            # Extract each image and copy it to the output folder with sequential naming
            for img_path in image_files:
                # Extract to temp directory
                docx_zip.extract(img_path, temp_dir)
                
                # Get image file extension
                _, ext = os.path.splitext(img_path)
                
                # Create sequentially named target path
                image_count += 1
                target_path = os.path.join(output_folder, f"image_{image_count:03d}{ext}")
                
                # Copy the file to the output folder
                shutil.copy2(os.path.join(temp_dir, img_path), target_path)
                print(f"Extracted: {target_path}")
    
    return image_count


def process_docx_or_folder(path):
    """
    Process a given path:
    1. If it's a folder, recursively process all files
    2. If it's a docx file, extract all images from it
    """
    if os.path.isdir(path):
        # It's a folder - process all files within recursively
        images_extracted = 0
        processed_files = 0
        
        for root, _, files in os.walk(path):
            for file_name in files:
                if file_name.lower().endswith('.docx'):
                    file_path = os.path.join(root, file_name)
                    # Create output folder with the same name as the docx (minus extension)
                    output_folder = os.path.join(root, os.path.splitext(file_name)[0])
                    
                    # Extract images from the docx
                    count = extract_images_from_docx(file_path, output_folder)
                    
                    if count > 0:
                        images_extracted += count
                        processed_files += 1
                        print(f"Extracted {count} images from {file_path} to {output_folder}")
        
        return processed_files, images_extracted
    
    elif os.path.isfile(path) and path.lower().endswith('.docx'):
        # It's a docx file - extract images directly
        output_folder = os.path.join(os.path.dirname(path), os.path.splitext(os.path.basename(path))[0])
        count = extract_images_from_docx(path, output_folder)
        
        if count > 0:
            print(f"Extracted {count} images from {path} to {output_folder}")
            return 1, count
        return 0, 0
    
    else:
        # Neither a folder nor a docx file
        print(f"Error: {path} is not a valid directory or Word document.")
        return 0, 0


def stitch_images_in_folder(folder_path, output_image=None, horizontal=False):
    """
    Stitch all images in a folder into a single image.
    
    Args:
        folder_path: Path to the folder containing images
        output_image: Path to save the stitched image (if None, uses folder name)
        horizontal: If True, stitch images side by side; otherwise stack vertically
    
    Returns:
        Path to the stitched image if successful, None otherwise
    """
    # Find and sort images
    image_paths = find_and_sort_images(folder_path)
    
    if not image_paths:
        print(f"No images found in folder: {folder_path}")
        return None
    
    # Open all images
    images = [Image.open(img_path) for img_path in image_paths]
    
    # Calculate dimensions of the stitched image
    if horizontal:
        # Images side by side
        width = sum(img.width for img in images)
        height = max(img.height for img in images)
    else:
        # Images stacked vertically
        width = max(img.width for img in images)
        height = sum(img.height for img in images)
    
    # Create a new blank image
    stitched_img = Image.new('RGB', (width, height), color=(255, 255, 255))
    
    # Paste images into the stitched image
    offset = 0
    for img in images:
        if horizontal:
            # Place side by side
            stitched_img.paste(img, (offset, 0))
            offset += img.width
        else:
            # Stack vertically
            stitched_img.paste(img, (0, offset))
            offset += img.height
    
    # If no output path specified, create it in the folder
    if output_image is None:
        folder_name = os.path.basename(folder_path)
        output_image = os.path.join(folder_path, f"{folder_name}_stitched.png")
    
    # Save the stitched image
    stitched_img.save(output_image)
    print(f"Stitched image saved: {output_image}")
    
    return output_image


app = typer.Typer(help="PDF Image Extraction and Conversion Utility - Process PDFs and images in various ways")


@app.command(help="Extract pages from PDFs as individual image files")
def extract_images(
    folders: List[str] = typer.Argument(
        ..., help="Folders containing PDF files to process"
    )
):
    """Process PDFs in the specified folders and extract images from them."""
    for folder in folders:
        if os.path.isdir(folder):
            typer.echo(f"Scanning folder: {folder}")
            process_pdf_in_folder(folder)
        else:
            typer.echo(f"Warning: {folder} is not a valid directory.")


@app.command(help="Convert PDFs to Word documents with each page as an embedded image")
def pdf_to_word(
    folders: List[str] = typer.Argument(
        ..., help="Folders containing PDF files to process"
    ),
    output_folder: Optional[str] = typer.Option(
        None, "--output", "-o", help="Output folder for the Word documents"
    ),
):
    """Convert PDFs to Word documents with each page as an image."""
    documents_created = []

    for folder in folders:
        if os.path.isdir(folder):
            typer.echo(f"Scanning folder: {folder}")
            docs = process_pdf_to_word(folder, output_folder)
            documents_created.extend(docs)
        else:
            typer.echo(f"Warning: {folder} is not a valid directory.")

    if documents_created:
        typer.echo(f"Created {len(documents_created)} Word document(s):")
        for doc in documents_created:
            typer.echo(f"  - {doc}")
    else:
        typer.echo("No Word documents were created.")


@app.command(help="Combine all images in a folder into a single PDF document")
def images_to_pdf(
    folder: str = typer.Argument(..., help="Folder containing image files to process"),
    output_pdf: Optional[str] = typer.Option(
        None, "--output", "-o", help="Output PDF file path"
    ),
):
    """Combine all images in a folder into a single PDF file, sorted by filename."""
    if os.path.isdir(folder):
        typer.echo(f"Processing images in folder: {folder}")
        # If no output path specified, create PDF within the folder
        if output_pdf is None:
            output_pdf = os.path.join(folder, os.path.basename(folder) + ".pdf")
        pdf_path = process_images_to_pdf(folder, output_pdf)
        if pdf_path:
            typer.echo(f"Created PDF: {pdf_path}")
    else:
        typer.echo(f"Error: {folder} is not a valid directory.")


@app.command(help="Convert PDFs to image-based PDFs - to prevent copying signatures")
def pdf_reimage(
    file_path: str = typer.Argument(..., help="PDF file or folder containing PDF files to process"),
):
    """
    For each PDF in the folder or the specified PDF file:
    1. Convert all pages to images
    2. Assemble those images back into a new PDF with '-images' suffix
    """
    converted_pdfs = []
    
    def process_single_pdf(pdf_path):
        """Helper function to process a single PDF file if valid"""
        if pdf_path.lower().endswith(".pdf") and not pdf_path.lower().endswith("-images.pdf"):
            typer.echo(f"Processing: {pdf_path}")
            new_pdf = convert_pdf_to_image_based_pdf(pdf_path)
            if new_pdf:
                converted_pdfs.append((pdf_path, new_pdf))
            return True
        return False
    
    if os.path.isdir(file_path):
        # Process all PDFs in the directory
        typer.echo(f"Processing all PDFs in directory: {file_path}")
        
        # Find all PDFs in the folder
        for root, _, files in os.walk(file_path):
            for file_name in files:
                pdf_path = os.path.join(root, file_name)
                process_single_pdf(pdf_path)
                
    elif os.path.isfile(file_path):
        # Process a single PDF file
        if not process_single_pdf(file_path):
            typer.echo(f"Error: {file_path} is not a PDF file.")
    else:
        typer.echo(f"Error: {file_path} does not exist or is not accessible.")
        return
    
    # Output summary
    if converted_pdfs:
        typer.echo(f"\nConverted {len(converted_pdfs)} PDF(s):")
        for original, new_pdf in converted_pdfs:
            typer.echo(f"  - {original} → {os.path.basename(new_pdf)}")
    else:
        typer.echo("No PDFs were converted.")


@app.command(help="Extract images from Word documents (.docx files)")
def docx_images(
    path: str = typer.Argument(..., help="Word document (.docx) or folder containing .docx files")
):
    """
    Extract all images from Word documents:
    1. If a folder is provided, process all .docx files recursively
    2. If a .docx file is provided, process just that file
    3. Images are saved in a folder with the same name as the document
    """
    processed_files, images_extracted = process_docx_or_folder(path)
    
    # Output summary
    if processed_files > 0:
        typer.echo(f"\nProcessed {processed_files} Word document(s), extracted {images_extracted} images.")
    else:
        typer.echo("No images were extracted from any Word documents.")


@app.command(help="Stitch all images in a folder into a single image")
def stitch_images(
    folder: str = typer.Argument(..., help="Folder containing image files to stitch"),
    output_image: Optional[str] = typer.Option(
        None, "--output", "-o", help="Output image file path"
    ),
    horizontal: bool = typer.Option(
        False, "--horizontal", "-h", help="Stitch images horizontally instead of vertically"
    )
):
    """
    Combine all images in a folder into a single stitched image.
    By default, images are stacked vertically in alphabetical order.
    """
    if os.path.isdir(folder):
        typer.echo(f"Stitching images in folder: {folder}")
        result = stitch_images_in_folder(folder, output_image, horizontal)
        if result:
            typer.echo(f"Created stitched image: {result}")
    else:
        typer.echo(f"Error: {folder} is not a valid directory.")


if __name__ == "__main__":
    app()
